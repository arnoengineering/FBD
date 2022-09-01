from docx.shared import Inches
from docx.shared import Pt
from docx import Document


class DateRange:  # todo add row col, space before
    def __init__(self, start, end, com="-"):
        self.start = start
        self.com = com
        self.end = end
        self.font = {"italics": True}
        # "Arial"
        # self.font_color = "Black"
        # self.font_style =
        self.font_size = 10
        self.page_width = 3.5

    def insert(self, para):
        par_f = para.paragraph_format
        par_f.add_tab_stop(Inches(self.page_width))
        p = para.add_run(self.start + self.com + self.end)
        f = p.font
        for i, v in self.font.items():
            setattr(f, i, v)


class SectionTitle:  # todo add insert
    def __init__(self, title):
        self.title = title.upper()

        self.font_color = "Blue"
        self.font_style = "Bold"

    def insert(self):
        doc.add_heading(self.title)


class Local:
    def __init__(self, com, loc="", job_title=""):
        self.com = com
        self.loc = loc  # todo same local dif font
        self.job_title = job_title

        self.split = "—"
        self.com_font = {
            "font_style": "Bold",
            "font_size": 14,
            "page_loc_x": 0.1
        }
        self.loc_font = {
            "font": "Arial",
            "font_color": "Blue",
            "font_style": "Bold",
            "font_size": 14,
            "page_loc_x": 0.1
        }

    def insert(self):
        p = doc.add_paragraph(self.com + ',')
        p.bold = True
        p.add_run(self.loc).bold = False
        p.add_run(self.loc + self.split).italics = True
        return p


class Experience:  # same as projects
    def __init__(self, com, start, end, info, loc="", job_title=""):
        self.head = Local(com, loc, job_title)
        self.date = DateRange(start, end)
        self.info = info
        self.page_loc_y = 10
        self.info_font = {

            "font": "Arial",
            "font_color": "Blue",
            "font_style": "Bold",
            "font_size": 14,
            "page_loc_x": 0.1,
            "inetnal_spacing": 1
        }
        # todo add row for x in info
        # todo add para sp

    def insert(self):
        hp = self.head.insert()
        self.date.insert(hp)
        for x in self.info:
            doc.add_paragraph(x, style='List Bullet 2')


class Volunteering:  # same as aditianal, but with date
    def __init__(self, dict_vol, date=None, title="VOLUNTEERING AND INTERESTS"):  # name local date
        self.title = SectionTitle(title)
        self.namefont = ""
        self.infofont = ""
        self.dict_vol = dict_vol  # todo add dates
        self.date = date

    def insert(self):
        self.title.insert()
        for x, y in self.dict_vol.items():
            p = doc.add_paragraph(x + ": ")
            p.bold = True
            p.add_run(y).italics = True  # rem bold
            if self.date:
                d = DateRange(*self.date)  # todo fix
                d.insert(p)


class CoOp:
    """
    UNIVERSITY OF ALBERTA - ACADEMIC & CO-OP STATUS
Mechanical Engineering, BSc Co-op, Class of 2024
Cumulative GPA					2.7/4.0
Completed Academic Terms				5 of 8
Completed Co-op Work Terms				1 of 5
Availability starting January 2022			 4-8 months
    """

    def __init__(self, gpa, terms, avail):
        self.title = SectionTitle("UNIVERSITY OF ALBERTA - ACADEMIC & CO-OP STATUS")
        self.name = Local("Mechanical Engineering, BSc Co-op", "Class of 2024")
        self.avail = avail
        self.info = {
            "Cumulative GPA": [gpa, "4.0", "/"],
            "Completed Academic Terms": [terms[0], "8", "of"],
            "Completed Co-op Work Terms": [terms[1], "5", "of"],
            f"Availability starting{avail}": [avail, "8 months"]
        }

    def insert(self):
        self.title.insert()
        for x, y in self.info.items():
            p = doc.add_paragraph(x)
            d = DateRange(*y)
            d.insert(p)


class References:
    def __init__(self):
        self.title = SectionTitle("REFERENCES")
        self.text = "Available upon request"

    def insert(self):
        self.title.insert()
        doc.add_paragraph(self.text)


class Header:  # todo head 2
    def __init__(self, name, email, web, image, im_w=150, im_h=70, adrress=None, number=None):
        self.head = doc.sections[0].header
        self.image = [image, im_w, im_h]
        self.name = name
        self.web = web
        self.adrress = adrress
        self.number = number
        if number:
            self.text = number + " | " + email
        else:
            self.text = email

    def insert(self):
        self.head.add_paragraph(self.name).bold = True  # todo style
        if self.adrress:
            self.head.add_paragraph(self.adrress)
        self.head.add_paragraph(self.text)
        self.head.add_picture(*self.image)


def save_doc(fname):
    doc.save(fname)
# todo end line per or none


doc = Document()
